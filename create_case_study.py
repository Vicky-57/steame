# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_steamee_case_study():
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styling helper
    def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    # Color definitions
    NAVY = RGBColor(27, 54, 93)     # #1B365D
    TEAL = RGBColor(13, 115, 119)   # #0D7377
    DARK_GRAY = RGBColor(74, 85, 104) # #4A5568
    BLACK = RGBColor(0, 0, 0)

    # 1. Header Block
    p_cs = doc.add_paragraph()
    p_cs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cs = p_cs.add_run("CASE STUDY")
    set_font(run_cs, size=12, bold=True, color=TEAL)
    p_cs.paragraph_format.space_after = Pt(2)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("STEAMEE")
    set_font(run_title, size=26, bold=True, color=NAVY)
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Data-Driven Operations & Marketing Analytics Platform")
    set_font(run_sub, size=14, bold=False, italic=True, color=DARK_GRAY)
    p_sub.paragraph_format.space_after = Pt(18)

    # 2. Main Hook
    p_hook = doc.add_paragraph()
    p_hook.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hook = p_hook.add_run("Steamee: Transforming D2C Garment Care with Strategic BI & Marketing Automation")
    set_font(run_hook, size=16, bold=True, color=NAVY)
    p_hook.paragraph_format.space_after = Pt(12)

    p_hook_desc = doc.add_paragraph()
    p_hook_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hook_desc = p_hook_desc.add_run(
        "How we architected a unified BI dashboard specification, solved duplicate KPI metrics, "
        "and mapped user event tracking for MoEngage integration — establishing a role-based operational "
        "and marketing command center."
    )
    set_font(run_hook_desc, size=11.5, italic=True, color=DARK_GRAY)
    p_hook_desc.paragraph_format.space_after = Pt(24)

    # 3. About the Client
    h_about = doc.add_heading(level=1)
    run_ha = h_about.add_run("About the Client")
    set_font(run_ha, size=18, bold=True, color=NAVY)
    h_about.paragraph_format.space_before = Pt(18)
    h_about.paragraph_format.space_after = Pt(6)

    p_about = doc.add_paragraph()
    run_pa = p_about.add_run(
        "Steamee is an emerging, tech-enabled D2C ironing and garment care ecosystem. Operating across a network of "
        "19 retail stores and localized ironing workshops, Steamee manages last-mile pickups and deliveries using "
        "their proprietary mobile application, WhatsApp channels, and call centers. Their mission is to modernize and structure "
        "the highly unorganized ironing industry through standardized workflows, doorstep convenience, and a scalable, "
        "franchise-friendly operational model."
    )
    set_font(run_pa, size=11)
    p_about.paragraph_format.space_after = Pt(12)

    # 4. The Challenge
    h_challenge = doc.add_heading(level=1)
    run_hc = h_challenge.add_run("The Challenge")
    set_font(run_hc, size=18, bold=True, color=NAVY)
    h_challenge.paragraph_format.space_before = Pt(18)
    h_challenge.paragraph_format.space_after = Pt(6)

    p_ch_intro = doc.add_paragraph()
    run_pci = p_ch_intro.add_run(
        "With thousands of recurring customers and dynamic daily order streams, Steamee faced two major operational "
        "and analytical challenges that limited their scalability:"
    )
    set_font(run_pci, size=11)
    p_ch_intro.paragraph_format.space_after = Pt(6)

    # Subheading 1: Operational Pain Points
    p_sh1 = doc.add_paragraph()
    run_sh1 = p_sh1.add_run("Operational & Customer Funnel Gaps")
    set_font(run_sh1, size=13, bold=True, color=TEAL)
    p_sh1.paragraph_format.space_before = Pt(8)
    p_sh1.paragraph_format.space_after = Pt(4)

    cx_bullets = [
        "Doorstep logistics were unoptimized, with pickup and delivery delays causing customer friction.",
        "No structured mechanism existed to measure individual pressman productivity, attendance shifts, or rider utilization.",
        "The customer onboarding journey suffered from undocumented drop-offs, making it hard to target dormant or churned users.",
        "The brand lacked a single tool to track and push marketing campaigns to inactive segments."
    ]
    for b in cx_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        run_b = p_b.add_run(b)
        set_font(run_b, size=11)
        p_b.paragraph_format.space_after = Pt(3)

    # Subheading 2: Business & Data Pain Points
    p_sh2 = doc.add_paragraph()
    run_sh2 = p_sh2.add_run("Business & Database Gaps")
    set_font(run_sh2, size=13, bold=True, color=TEAL)
    p_sh2.paragraph_format.space_before = Pt(8)
    p_sh2.paragraph_format.space_after = Pt(4)

    biz_bullets = [
        "The existing admin portal was plagued by duplicate KPIs across multiple tabs, yielding conflicting totals for sales and user segments.",
        "Critical portal software bugs caused negative order counts in the analytics APIs and crashed homepage date filters.",
        "Leadership had zero visibility into prepaid wallet outstanding liabilities, unutilized cash floats, and expired reward coins.",
        "No marketing spend tables or ad account APIs were connected to calculate blended or channel-specific CAC and ROAS."
    ]
    for b in biz_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        run_b = p_b.add_run(b)
        set_font(run_b, size=11)
        p_b.paragraph_format.space_after = Pt(3)

    # 5. Our Solution
    h_sol = doc.add_heading(level=1)
    run_hsol = h_sol.add_run("Our Solution: Unified BI & Marketing Blueprint")
    set_font(run_hsol, size=18, bold=True, color=NAVY)
    h_sol.paragraph_format.space_before = Pt(18)
    h_sol.paragraph_format.space_after = Pt(6)

    p_sol_intro = doc.add_paragraph()
    run_psi = p_sol_intro.add_run(
        "We designed and delivered an enterprise-grade analytics specification blueprint and integration mapping "
        "to resolve data clutter and establish a robust, role-based marketing and decision command center."
    )
    set_font(run_psi, size=11)
    p_sol_intro.paragraph_format.space_after = Pt(12)

    # Solution Components
    sol_components = [
        ("Unified Specification Sheets (13 Tabs)", 
         "We structured a clean, page-by-page frontend UI spec detailing all 11 core dashboards. "
         "Every visual element, headline card, and grid is mapped with its Target Role, Available Filters, and Suggested Visualization. "
         "This establishes an objective design manual for the frontend developers."),
         
        ("MoEngage Event Schema & Attributes", 
         "We mapped a comprehensive custom user event schema covering the 4 lifecycle stages: Acquisition (sign_up, drop_off_screen), "
         "Activation (add_to_cart, order_completed), Retention (feedback_submitted, bags_linked), and Referrals. "
         "We also standardized custom user profile attributes to create dynamic customer segments."),
         
        ("Dual-Architecture Strategy (Sheets 12 & 13)", 
         "We provided two separate, parallel implementation pathways: "
         "(1) Option 1: Portal + MoEngage for custom in-house dashboard screens and backend endpoints, and "
         "(2) Option 2: Metabase + MoEngage for fast SQL-based panels on top of the Postgres database. "
         "This allows leadership to select the ideal speed-to-value implementation path."),
         
        ("API Gap Analysis & Reconciliations", 
         "We audited 130+ metric requests and mapped them to database fields. We consolidated 8 groups of duplicate metrics "
         "into single definitions, and identified the exact 7 backend APIs needed to plug operational data gaps (e.g. stock inventory, "
         "machine maintenance, and ad account spend).")
    ]

    for title, desc in sol_components:
        p_comp = doc.add_paragraph()
        run_ct = p_comp.add_run(f"{title}: ")
        set_font(run_ct, size=11.5, bold=True, color=TEAL)
        run_cd = p_comp.add_run(desc)
        set_font(run_cd, size=11)
        p_comp.paragraph_format.space_after = Pt(6)

    # 6. Technology Highlights (Stats callouts style)
    h_tech = doc.add_heading(level=1)
    run_htech = h_tech.add_run("Project Highlights")
    set_font(run_htech, size=18, bold=True, color=NAVY)
    h_tech.paragraph_format.space_before = Pt(18)
    h_tech.paragraph_format.space_after = Pt(6)

    tech_stats = [
        "13 Dashboard Spec Sheets",
        "59 Mapped KPIs & SQL Formulas",
        "15 Configuration Master Tables Documented",
        "2 Architectural Options (Portal vs. Metabase)",
        "4 Lifecycle Funnel Stages Mapped for MoEngage"
    ]
    for stat in tech_stats:
        p_s = doc.add_paragraph(style='List Bullet')
        run_s = p_s.add_run(stat)
        set_font(run_s, size=11.5, bold=True, color=TEAL)
        p_s.paragraph_format.space_after = Pt(4)

    # 7. Results & Impact
    h_results = doc.add_heading(level=1)
    run_hres = h_results.add_run("Results & Impact")
    set_font(run_hres, size=18, bold=True, color=NAVY)
    h_results.paragraph_format.space_before = Pt(18)
    h_results.paragraph_format.space_after = Pt(6)

    results_bullets = [
        "100% Alignment on Metric Definitions: Consolidated all duplicate views (sales, AOV, active segments) into single canonical definitions locked in the database.",
        "Zero-Blindspot Operations: Identified and flagged 4 major portal bugs and 4 critical data gaps before dashboard construction began.",
        "Faster Dashboard Rollout: Mapped and categorized all deliverables into P0, P1, and P2 priorities, enabling developers to build high-priority tiles in under one sprint.",
        "Automated Marketing Activation: Reconciled customer lifecycle recency windows with MoEngage triggers, enabling automated push campaigns for dormant users.",
        "Unified Governance: Established clear data health, freshness, and ingestion quality rules to maintain reports integrity over time."
    ]
    for rb in results_bullets:
        p_r = doc.add_paragraph(style='List Bullet')
        run_r = p_r.add_run(rb)
        set_font(run_r, size=11)
        p_r.paragraph_format.space_after = Pt(4)

    # 8. Testimonial
    h_test = doc.add_heading(level=1)
    run_htest = h_test.add_run("Client Testimonial")
    set_font(run_htest, size=18, bold=True, color=NAVY)
    h_test.paragraph_format.space_before = Pt(18)
    h_test.paragraph_format.space_after = Pt(6)

    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.5)
    run_t = p_t.add_run(
        "\"The analytics blueprint and MoEngage integration specs were exactly what our engineering and marketing "
        "teams needed. It resolved our duplicate metrics, gave us a clear division of work between MoEngage and our internal portal, "
        "and saved us months of development planning. We now have a clear roadmap to measure store performance and track our "
        "500 Super Value customer targets in real-time. This has been a game-changer for Steamee's data strategy.\""
    )
    set_font(run_t, size=11, italic=True, color=DARK_GRAY)
    p_t.paragraph_format.space_after = Pt(4)

    p_ta = doc.add_paragraph()
    p_ta.paragraph_format.left_indent = Inches(0.5)
    run_ta = p_ta.add_run("— Founder & Leadership Team, Steamee")
    set_font(run_ta, size=10, bold=True, color=TEAL)
    p_ta.paragraph_format.space_after = Pt(18)

    # 9. Why It Works
    h_why = doc.add_heading(level=1)
    run_hwhy = h_why.add_run("Why It Works")
    set_font(run_hwhy, size=18, bold=True, color=NAVY)
    h_why.paragraph_format.space_before = Pt(18)
    h_why.paragraph_format.space_after = Pt(6)

    p_why = doc.add_paragraph()
    run_pw = p_why.add_run(
        "The Steamee data project succeeds because it bridges the communication gap between business goals and technical "
        "execution. By delivering a highly structured, role-based UI blueprint and exact SQL formulas before any coding began, "
        "we eliminated the developer guesswork that typically results in broken filters and incorrect aggregations. "
        "Furthermore, by analyzing and mapping MoEngage events alongside internal database queries, we provided the client "
        "with an integrated operational AND marketing system, instead of a static, un-actionable report grid."
    )
    set_font(run_pw, size=11)
    p_why.paragraph_format.space_after = Pt(12)

    # Save document
    doc.save("STEAMEE_Case_Study.docx")
    print("STEAMEE_Case_Study.docx generated successfully!")

if __name__ == "__main__":
    create_steamee_case_study()
